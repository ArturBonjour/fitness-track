#!/usr/bin/env python3
"""
convert_diploma.py — конвертация дипломной работы из Markdown в DOCX.

Требования:
    pip install python-docx

Использование:
    python3 convert_diploma.py                        # создаёт diploma.docx рядом со скриптом
    python3 convert_diploma.py -o /path/to/output.docx

Форматирование соответствует требованиям методических указаний КФУ (КАДиТП, 2025):
  - Шрифт Times New Roman 14 пт, полуторный (1.5×) интервал
  - Поля: верхнее 20 мм, нижнее 20 мм, правое 15 мм, левое 30 мм
  - Абзацный отступ основного текста 1.25 см
  - Заголовки по центру, без жирного, без курсива, без точки в конце
  - Каждая глава начинается с новой страницы
  - Нумерация страниц снизу по центру, Times New Roman 12 пт
  - Блоки кода — Courier New 12 пт, рамка 0.25 пт (2/8 pt), белый фон
  - Подписи к таблицам выровнены по правому краю
  - Таблицы растянуты на всю ширину текстового поля
"""

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SOURCE_MD = Path(__file__).parent / "diploma_ch1_draft.md"
DEFAULT_OUT = Path(__file__).parent / "diploma.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE_MAIN = 14          # пт — основной текст
FONT_SIZE_HEADING = 14       # пт — заголовки (методичка: такой же размер)
FONT_SIZE_CODE = 12          # пт — блоки кода
FONT_SIZE_PAGE_NUM = 12      # пт — номер страницы

MARGIN_TOP    = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT   = Cm(3.0)
MARGIN_RIGHT  = Cm(1.5)

PARA_INDENT   = Cm(1.25)     # абзацный отступ основного текста

# Главы, которые должны начинаться с новой страницы (уровень 1)
# Подглавы (уровень 2) — не начинаются с новой страницы
CHAPTER_NEW_PAGE = True


# ---------------------------------------------------------------------------
# Helpers — низкоуровневые операции над XML OOXML
# ---------------------------------------------------------------------------

def set_page_margins(doc: Document) -> None:
    """Задаёт поля страницы согласно методическим указаниям."""
    section = doc.sections[0]
    section.top_margin    = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin   = MARGIN_LEFT
    section.right_margin  = MARGIN_RIGHT


def add_page_numbers(doc: Document) -> None:
    """Добавляет сквозную нумерацию страниц снизу по центру (12 пт, TNR).

    Нумерация начинается с титульного листа, но номер на нём не отображается;
    первый видимый номер — на странице «Содержание». Реализуется через
    «different first page footer»: первая страница имеет пустой колонтитул,
    все остальные — с полем PAGE.
    """
    section = doc.sections[0]
    # Титульный лист: отдельный (пустой) колонтитул первой страницы
    section.different_first_page_header_footer = True
    # Обычный колонтитул (со 2-й страницы и далее)
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = para.add_run()
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PAGE_NUM)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # Первая страница: явно оставляем пустой колонтитул (nothing to add)
    _ = section.first_page_footer


def _set_line_spacing_15(fmt) -> None:
    """Устанавливает полуторный (1.5×) межстрочный интервал (правило MULTIPLE).

    В python-docx: если передать float, интервал интерпретируется как кратный
    (MULTIPLE), если передать Length — как точный (EXACTLY). Поэтому используем 1.5.
    """
    fmt.line_spacing = 1.5   # float → WD_LINE_SPACING.MULTIPLE автоматически


def _set_paragraph_format(
    para,
    indent_first: bool = True,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
) -> None:
    """Применяет единые настройки абзаца для основного текста.

    Включает контроль висячих строк (widowControl) — методичка запрещает
    менее 5 строк на странице.
    """
    fmt = para.paragraph_format
    fmt.alignment = alignment
    fmt.space_after = Pt(0)
    _set_line_spacing_15(fmt)
    fmt.first_line_indent = PARA_INDENT if indent_first else Pt(0)
    _set_widow_orphan_control(para)


def _apply_font(run, size: int = FONT_SIZE_MAIN, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def _add_page_break_before(para) -> None:
    """Вставляет разрыв страницы перед абзацем (w:pageBreakBefore)."""
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement("w:pageBreakBefore")
    pPr.append(pb)


def _set_keep_with_next(para) -> None:
    """Запрещает отрыв заголовка от следующего абзаца (w:keepNext).

    Методичка: «заголовки на странице не должны отрываться от текста,
    то есть название не должно быть на одной странице, а текст — на другой.
    После названия должно быть не менее 3 строк текста».
    keepNext в сочетании с keepLines обеспечивает это требование.
    """
    pPr = para._p.get_or_add_pPr()
    kn = OxmlElement("w:keepNext")
    pPr.append(kn)
    kl = OxmlElement("w:keepLines")
    pPr.append(kl)


def _set_widow_orphan_control(para) -> None:
    """Включает контроль висячих строк (w:widowControl).

    Методичка: «на странице не допускается менее 5 строк».
    widowControl предотвращает появление одиночных строк абзаца
    в начале или конце страницы.
    """
    pPr = para._p.get_or_add_pPr()
    wc = OxmlElement("w:widowControl")
    wc.set(qn("w:val"), "1")
    pPr.append(wc)


def _set_table_full_width(table) -> None:
    """Растягивает таблицу на всю ширину текстового поля."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _add_code_border(para) -> None:
    """Добавляет рамку 0.25 пт вокруг абзаца с кодом (имитация рисунка с кодом)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        bdr = OxmlElement(f"w:{side}")
        bdr.set(qn("w:val"), "single")
        bdr.set(qn("w:sz"), "2")      # 2 восьмых пункта = 0.25 пт (по методичке)
        bdr.set(qn("w:space"), "4")
        bdr.set(qn("w:color"), "000000")
        pBdr.append(bdr)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Парсер Markdown → Document
# ---------------------------------------------------------------------------

class DiplomaMdParser:
    """
    Упрощённый парсер подмножества Markdown, используемого в дипломе.

    Поддерживает: заголовки #/##, параграфы, нумерованные и маркированные
    списки, таблицы, блоки кода, горизонтальные разделители.
    """

    HEADING_RE   = re.compile(r"^(#{1,6})\s+(.*)")
    HR_RE        = re.compile(r"^-{3,}$")
    ORDERED_RE   = re.compile(r"^(\d+)\.\s+(.*)")
    UNORDERED_RE = re.compile(r"^[-–]\s+(.*)")
    TABLE_ROW_RE = re.compile(r"^\|")
    TABLE_CAP_RE = re.compile(r"^Таблица\s+[\dА]")
    FIGURE_CAP_RE = re.compile(r"^Рисунок\s+\d")
    CODE_FENCE   = "```"

    # Структурные заголовки, которые начинают новую страницу (уровень 1)
    TOP_LEVEL_WORDS = {
        "ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "СОДЕРЖАНИЕ",
    }

    def __init__(self, doc: Document):
        self.doc = doc
        self._in_code = False
        self._code_lang = ""
        self._code_lines: list[str] = []
        self._table_rows: list[list[str]] = []
        self._in_table = False
        self._first_chapter = True   # первая глава не нуждается в разрыве (идёт первой)

    def parse(self, text: str) -> None:
        lines = text.splitlines()
        i = 0
        while i < len(lines):
            line = lines[i]

            # --- code fence ---
            if line.startswith(self.CODE_FENCE):
                if not self._in_code:
                    self._flush_table()
                    self._in_code = True
                    self._code_lang = line[len(self.CODE_FENCE):].strip()
                    self._code_lines = []
                else:
                    self._flush_code()
                    self._in_code = False
                i += 1
                continue

            if self._in_code:
                self._code_lines.append(line)
                i += 1
                continue

            # --- horizontal rule (title page decorators — skip) ---
            if self.HR_RE.match(line):
                i += 1
                continue

            # --- table row ---
            if self.TABLE_ROW_RE.match(line):
                self._in_table = True
                cells = [c.strip() for c in line.strip("|").split("|")]
                # Skip Markdown separator rows (|---|---|)
                if all(re.match(r"^[-: ]+$", c) for c in cells):
                    i += 1
                    continue
                self._table_rows.append(cells)
                i += 1
                continue
            else:
                if self._in_table:
                    self._flush_table()

            # --- heading ---
            m = self.HEADING_RE.match(line)
            if m:
                level = len(m.group(1))
                text_content = m.group(2).strip()
                self._add_heading(text_content, level)
                i += 1
                continue

            # --- ordered list item ---
            m = self.ORDERED_RE.match(line)
            if m:
                self._add_list_item(m.group(1) + ". " + m.group(2), ordered=True)
                i += 1
                continue

            # --- unordered list item ---
            m = self.UNORDERED_RE.match(line)
            if m:
                self._add_list_item(m.group(1), ordered=False)
                i += 1
                continue

            # --- empty line ---
            if not line.strip():
                i += 1
                continue

            # --- regular paragraph (including table/figure captions) ---
            self._add_paragraph(line)
            i += 1

        self._flush_table()
        if self._in_code:
            self._flush_code()

    # -----------------------------------------------------------------------
    # Private rendering methods
    # -----------------------------------------------------------------------

    def _needs_page_break(self, text: str, level: int) -> bool:
        """Возвращает True, если перед заголовком нужен разрыв страницы."""
        if level != 1:
            return False
        # Уровень 1: всегда новая страница, кроме первого элемента документа
        if self._first_chapter:
            self._first_chapter = False
            return False
        return True

    def _add_heading(self, text: str, level: int) -> None:
        """Добавляет заголовок главы (уровень 1) или подглавы (уровень 2).

        Методичка: заголовки без жирного, без курсива, без подчёркивания,
        выравнивание по центру, без точки в конце, без абзацного отступа.
        """
        # Убираем Markdown-маркеры жирного (из TOC)
        text = text.replace("**", "")

        need_break = self._needs_page_break(text, level)

        para = self.doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Pt(0)
        # Перед подглавой — одна пустая строка (12 пт), перед главой — 0
        fmt.space_before = Pt(12) if level > 1 else Pt(0)
        fmt.space_after  = Pt(12)
        _set_line_spacing_15(fmt)

        if need_break:
            _add_page_break_before(para)

        # Не допускаем отрыв заголовка от следующего текста (методичка §2.2)
        _set_keep_with_next(para)

        run = para.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_HEADING)
        run.font.bold   = False
        run.font.italic = False
        run.font.underline = False

        # Назначаем стиль Word, чтобы автооглавление могло подхватить заголовок
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
        if level in style_map:
            try:
                para.style = self.doc.styles[style_map[level]]
                # Стиль может перезаписать форматирование — возвращаем его явно
                fmt2 = para.paragraph_format
                fmt2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fmt2.first_line_indent = Pt(0)
                fmt2.space_before = Pt(12) if level > 1 else Pt(0)
                fmt2.space_after  = Pt(12)
                _set_line_spacing_15(fmt2)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_HEADING)
                run.font.bold   = False
                run.font.italic = False
            except KeyError:
                pass

    def _add_paragraph(self, text: str) -> None:
        """Добавляет параграф основного текста с разбором inline-разметки."""
        # Подпись к рисунку: "Рисунок X.Y. ..."
        is_fig_caption = bool(self.FIGURE_CAP_RE.match(text))
        # Подпись к таблице: "Таблица X.X — ..."
        is_tbl_caption = bool(self.TABLE_CAP_RE.match(text))

        if is_tbl_caption:
            # Методичка: подпись к таблице выравнивается по правому краю
            para = self.doc.add_paragraph()
            fmt = para.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fmt.first_line_indent = Pt(0)
            fmt.space_before = Pt(12)
            fmt.space_after  = Pt(0)
            _set_line_spacing_15(fmt)
            self._add_inline_text(para, text)
            return

        if is_fig_caption:
            # Подпись к рисунку: по центру, пустая строка после
            para = self.doc.add_paragraph()
            fmt = para.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = Pt(0)
            fmt.space_before = Pt(0)
            fmt.space_after  = Pt(12)
            _set_line_spacing_15(fmt)
            self._add_inline_text(para, text)
            return

        # Обычный абзац
        para = self.doc.add_paragraph()
        _set_paragraph_format(para)
        self._add_inline_text(para, text)

    def _add_list_item(self, text: str, ordered: bool) -> None:
        """Добавляет элемент нумерованного или маркированного списка.

        Методичка: маркер — тире «–», нумерованные — арабские цифры.
        Отступ слева равен абзацному отступу основного текста (1.25 см).
        """
        para = self.doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Pt(0)
        fmt.left_indent = PARA_INDENT
        fmt.space_after  = Pt(0)
        _set_line_spacing_15(fmt)

        prefix = "– " if not ordered else ""
        self._add_inline_text(para, prefix + text)

    def _add_inline_text(self, para, text: str) -> None:
        """Разбирает inline **bold** и `code`, добавляет runs."""
        token_re = re.compile(r"(\*\*.*?\*\*|`[^`]+`)")
        parts = token_re.split(text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                content = part[2:-2]
                run = para.add_run(content)
                # Методичка запрещает жирный в основном тексте
                _apply_font(run, bold=False)
            elif part.startswith("`") and part.endswith("`"):
                content = part[1:-1]
                run = para.add_run(content)
                run.font.name = "Courier New"
                run.font.size = Pt(FONT_SIZE_CODE)
            else:
                run = para.add_run(part)
                _apply_font(run)

    def _flush_code(self) -> None:
        """Добавляет блок кода как абзац Courier New с рамкой 0.25 пт.

        Методичка: фрагменты кода ≤ ½ страницы оформляются как рисунки
        (белый фон, рамка чёрного цвета 0.25 пт, шрифт 10–14 пт).
        """
        if not self._code_lines:
            self._code_lang = ""
            return

        para = self.doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt.first_line_indent = Pt(0)
        fmt.left_indent  = Pt(0)
        fmt.right_indent = Pt(0)
        fmt.space_before = Pt(6)
        fmt.space_after  = Pt(6)
        # Одинарный интервал для кода (читабельнее в моноширинном шрифте)
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE

        for idx, code_line in enumerate(self._code_lines):
            run = para.add_run(code_line)
            run.font.name = "Courier New"
            run.font.size = Pt(FONT_SIZE_CODE)
            run.font.bold   = False
            run.font.italic = False
            if idx < len(self._code_lines) - 1:
                run.add_break()

        _add_code_border(para)

        self._code_lines = []
        self._code_lang = ""

    def _flush_table(self) -> None:
        """Добавляет накопленные строки как Word-таблицу на полную ширину страницы."""
        if not self._table_rows:
            self._in_table = False
            return

        rows = self._table_rows
        col_count = max(len(r) for r in rows)

        table = self.doc.add_table(rows=len(rows), cols=col_count)
        table.style = "Table Grid"
        _set_table_full_width(table)

        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx]
            is_header = (r_idx == 0)
            for c_idx in range(col_count):
                cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                cell = row.cells[c_idx]
                cell.text = ""
                cell_text_clean = cell_text.replace("**", "")
                para = cell.paragraphs[0]
                # Заголовок таблицы — по центру, остальные строки — по ширине
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if is_header
                    else WD_ALIGN_PARAGRAPH.JUSTIFY
                )
                fmt = para.paragraph_format
                fmt.space_after = Pt(0)
                _set_line_spacing_15(fmt)
                run = para.add_run(cell_text_clean)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE_MAIN)
                run.font.bold = is_header

        self._table_rows = []
        self._in_table = False

        # Пустая строка после таблицы (методичка: «после таблицы до текста пустая строка»)
        spacer = self.doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def build_docx(source_md: Path, output_path: Path) -> None:
    doc = Document()

    # Настройка страницы
    set_page_margins(doc)
    add_page_numbers(doc)

    # Базовый стиль «Normal» — Times New Roman 14 пт
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(FONT_SIZE_MAIN)

    # Читаем и очищаем исходник
    text = source_md.read_text(encoding="utf-8")
    # Убираем YAML front-matter (--- ... ---), если есть
    text = re.sub(r"^---\s*\n.*?\n---\s*\n", "", text, count=1, flags=re.DOTALL)

    # Парсим и строим документ
    parser = DiplomaMdParser(doc)
    parser.parse(text)

    doc.save(str(output_path))
    print(f"[OK] Сохранено: {output_path}")
    print(f"     Абзацев в документе: {len(doc.paragraphs)}")


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Конвертация diploma_ch1_draft.md → DOCX (КФУ КАДиТП 2025)"
    )
    ap.add_argument(
        "-i", "--input",
        default=str(SOURCE_MD),
        help="Входной Markdown-файл (по умолчанию: diploma_ch1_draft.md)",
    )
    ap.add_argument(
        "-o", "--output",
        default=str(DEFAULT_OUT),
        help="Выходной DOCX-файл (по умолчанию: diploma.docx)",
    )
    args = ap.parse_args()

    src = Path(args.input)
    out = Path(args.output)

    if not src.exists():
        print(f"[ОШИБКА] Файл не найден: {src}", file=sys.stderr)
        sys.exit(1)

    out.parent.mkdir(parents=True, exist_ok=True)
    build_docx(src, out)


if __name__ == "__main__":
    main()
